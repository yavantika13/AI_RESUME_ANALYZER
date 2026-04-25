# import streamlit as st
# import pdfplumber
# from docx import Document

# # ------------------ Page Config ------------------
# st.set_page_config(page_title="AI Resume Analyzer", layout="centered")

# st.title("AI Resume Analyzer")
# st.write("Upload your resume (PDF or DOCX)")

# # ------------------ Helper Functions ------------------
# def read_pdf(file):
#     text = ""
#     with pdfplumber.open(file) as pdf:
#         for page in pdf.pages:
#             extracted = page.extract_text()
#             if extracted:
#                 text += extracted
#     return text

# def read_docx(file):
#     doc = Document(file)
#     return "\n".join([p.text for p in doc.paragraphs])

# def extract_skills(text):
#     skills_db = [
#         "python", "java", "c++", "html", "css", "javascript",
#         "react", "sql", "machine learning", "data analysis",
#         "power bi", "excel", "azure", "aws"
#     ]

#     found_skills = []
#     text_lower = text.lower()

#     for skill in skills_db:
#         if skill in text_lower:
#             found_skills.append(skill.capitalize())

#     return list(set(found_skills))

# def resume_score(skills):
#     return min(len(skills) * 10, 100)

# def recommended_role(skills):
#     if "React" in skills or "Javascript" in skills:
#         return "Frontend Developer"
#     elif "Python" in skills and "Machine learning" in skills:
#         return "AI / ML Engineer"
#     elif "Sql" in skills or "Power bi" in skills:
#         return "Data Analyst"
#     else:
#         return "Software Developer"

# # ------------------ File Upload ------------------
# uploaded_file = st.file_uploader(
#     "Choose a resume",
#     type=["pdf", "docx"]
# )

# if uploaded_file:
#     # Read resume
#     if uploaded_file.name.endswith(".pdf"):
#         resume_text = read_pdf(uploaded_file)
#     else:
#         resume_text = read_docx(uploaded_file)

#     # ------------------ Resume Text ------------------
#     st.subheader("Extracted Resume Text")
#     st.text_area("Resume Content", resume_text, height=300)

#     # ------------------ Skill Extraction ------------------
#     skills = extract_skills(resume_text)

#     st.subheader("Detected Skills")

#     if skills:
#         cols = st.columns(len(skills))
#         for col, skill in zip(cols, skills):
#             col.markdown(
#                 f"""
#                 <div style="
#                     background-color:#e0f2fe;
#                     padding:8px 12px;
#                     border-radius:20px;
#                     text-align:center;
#                     font-weight:600;
#                     color:#0369a1;
#                     ">
#                     {skill}
#                 </div>
#                 """,
#                 unsafe_allow_html=True
#             )
#     else:
#         st.warning("No skills detected")

#     # ------------------ Resume Score ------------------
#     score = resume_score(skills)

#     st.subheader("Resume Score")
#     st.progress(score / 100)
#     st.write(f"**{score} / 100**")
#     st.caption("💡 Improve your score by adding more relevant technical skills")

#     # ------------------ Suggestions ------------------
#     st.subheader("Suggestions to Improve Resume")

#     recommended_skills = ["Python", "SQL", "Azure", "Machine learning", "Power bi"]
#     missing_skills = [s for s in recommended_skills if s not in skills]

#     if missing_skills:
#         for skill in missing_skills:
#             st.write(f"• Add **{skill}** to strengthen your resume")
#     else:
#         st.success("Your resume already contains strong technical skills!")

#     # ------------------ Job Role ------------------
#     st.subheader("Recommended Job Role")
#     role = recommended_role(skills)
#     st.success(role)



import streamlit as st
import pdfplumber
from docx import Document
import plotly.express as px
import pandas as pd

# ------------------ Page Config ------------------
st.set_page_config(page_title="AI Resume Analyzer", layout="wide")

# ------------------ Custom CSS ------------------
st.markdown("""
<style>
.skill-box {
    background-color:#e0f2fe;
    padding:8px 12px;
    border-radius:20px;
    text-align:center;
    font-weight:600;
    color:#0369a1;
    margin:5px;
}
</style>
""", unsafe_allow_html=True)

# ------------------ Title ------------------
st.title("🚀 AI Resume Analyzer")
st.write("Upload your resume (PDF or DOCX)")

# ------------------ Helper Functions ------------------
def read_pdf(file):
    text = ""
    with pdfplumber.open(file) as pdf:
        for page in pdf.pages:
            extracted = page.extract_text()
            if extracted:
                text += extracted
    return text

def read_docx(file):
    doc = Document(file)
    return "\n".join([p.text for p in doc.paragraphs])

def extract_skills(text):
    skills_db = [
        "python", "java", "c++", "html", "css", "javascript",
        "react", "sql", "machine learning", "data analysis",
        "power bi", "excel", "azure", "aws"
    ]

    found_skills = []
    text_lower = text.lower()

    for skill in skills_db:
        if skill in text_lower:
            found_skills.append(skill.capitalize())

    return list(set(found_skills))

def resume_score(skills):
    return min(len(skills) * 10, 100)

def recommended_role(skills):
    if "React" in skills or "Javascript" in skills:
        return "Frontend Developer"
    elif "Python" in skills and "Machine learning" in skills:
        return "AI / ML Engineer"
    elif "Sql" in skills or "Power bi" in skills:
        return "Data Analyst"
    else:
        return "Software Developer"

# ------------------ File Upload ------------------
uploaded_file = st.file_uploader("📄 Choose a resume", type=["pdf", "docx"])

if uploaded_file:
    # Read resume
    if uploaded_file.name.endswith(".pdf"):
        resume_text = read_pdf(uploaded_file)
    else:
        resume_text = read_docx(uploaded_file)

    # Layout split
    col1, col2 = st.columns([1,1])

    # ------------------ LEFT SIDE ------------------
    with col1:
        st.subheader("📄 Extracted Resume Text")
        st.text_area("Resume Content", resume_text, height=400)

    # ------------------ RIGHT SIDE ------------------
    with col2:
        skills = extract_skills(resume_text)

        st.subheader("🧠 Detected Skills")

        if skills:
            skill_cols = st.columns(3)
            for i, skill in enumerate(skills):
                skill_cols[i % 3].markdown(f"<div class='skill-box'>{skill}</div>", unsafe_allow_html=True)
        else:
            st.warning("No skills detected")

        # ------------------ Score ------------------
        score = resume_score(skills)

        st.subheader("📊 Resume Score")
        st.progress(score / 100)
        st.write(f"### {score} / 100")

        # ------------------ Role ------------------
        st.subheader("🎯 Recommended Role")
        role = recommended_role(skills)
        st.success(role)

    # ================== 📊 INTERACTIVE CHARTS ==================
    if skills:
        st.subheader("📊 Skills Dashboard")

        col3, col4 = st.columns(2)

        # ---------- Bar Chart ----------
        with col3:
            df_skills = pd.DataFrame({
                "Skill": skills,
                "Count": [1] * len(skills)
            })

            fig_bar = px.bar(df_skills, x="Skill", y="Count",
                             title="Skills Detected",
                             text_auto=True)
            st.plotly_chart(fig_bar, use_container_width=True)

        # ---------- Pie Chart ----------
        with col4:
            tech_skills = ["Python", "Java", "C++", "React", "Sql", "Machine learning"]

            tech_count = len([s for s in skills if s in tech_skills])
            other_count = len(skills) - tech_count

            df_pie = pd.DataFrame({
                "Category": ["Technical", "Other"],
                "Value": [tech_count, other_count]
            })

            fig_pie = px.pie(df_pie, names="Category", values="Value",
                             title="Skills Distribution")
            st.plotly_chart(fig_pie, use_container_width=True)

    # ------------------ Score Breakdown ------------------
    st.subheader("📈 Resume Score Breakdown")

    skill_score = len(skills) * 10
    education_score = 20
    experience_score = 10

    df_score = pd.DataFrame({
        "Category": ["Skills", "Education", "Experience"],
        "Score": [skill_score, education_score, experience_score]
    })

    fig_score = px.bar(df_score, x="Category", y="Score",
                       title="Score Breakdown",
                       text_auto=True)

    st.plotly_chart(fig_score, use_container_width=True)

    # ------------------ Suggestions ------------------
    st.subheader("💡 Suggestions to Improve Resume")

    recommended_skills = ["Python", "SQL", "Azure", "Machine learning", "Power bi"]
    missing_skills = [s for s in recommended_skills if s not in skills]

    if missing_skills:
        for skill in missing_skills:
            st.write(f"• Add **{skill}** to strengthen your resume")
    else:
        st.success("Your resume already contains strong technical skills!")